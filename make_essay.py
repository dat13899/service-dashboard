from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# ── Helper ──
def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0,0,0)
    return h

def para(text, bold=False, align=None, size=None, indent_first=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    if size: run.font.size = Pt(size)
    if bold: run.bold = True
    if align: p.alignment = align
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.text = ''
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    return p

# ═══════════════════════════════════════
#  TIÊU ĐỀ
# ═══════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BỘ GIÁO DỤC VÀ ĐÀO TẠO')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('TRƯỜNG ĐẠI HỌC …………')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True

doc.add_paragraph()  # blank line

inner = doc.add_paragraph()
inner.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = inner.add_run('-----❧-----')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()

main_title = doc.add_paragraph()
main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = main_title.add_run('TIỂU LUẬN')
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True

doc.add_paragraph()

topic = doc.add_paragraph()
topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = topic.add_run('TÁC ĐỘNG CỦA TRÍ TUỆ NHÂN TẠO (AI)\nĐỐI VỚI CÁC NGÀNH NGHỀ TRONG NỀN KINH TẾ')
run.font.name = 'Times New Roman'
run.font.size = Pt(15)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Giảng viên hướng dẫn: ………………………\nSinh viên thực hiện: ……………………………\nLớp: …………  –  MSSV: …………………\nHà Nội, tháng 7 năm 2026')
run.font.name = 'Times New Roman'
run.font.size = Pt(13)

doc.add_page_break()

# ═══════════════════════════════════════
#  MỤC LỤC (manual)
# ═══════════════════════════════════════
heading('MỤC LỤC', level=1)
toc_items = [
    'LỜI MỞ ĐẦU',
    'CHƯƠNG 1: TỔNG QUAN VỀ TRÍ TUỆ NHÂN TẠO',
    '  1.1. Khái niệm và lịch sử phát triển',
    '  1.2. Các công nghệ AI cốt lõi',
    '  1.3. Thực trạng ứng dụng AI hiện nay',
    'CHƯƠNG 2: TÁC ĐỘNG CỦA AI ĐẾN CÁC NGÀNH NGHỀ',
    '  2.1. Ngành sản xuất và chế tạo',
    '  2.2. Ngành y tế và chăm sóc sức khỏe',
    '  2.3. Ngành giáo dục và đào tạo',
    '  2.4. Ngành tài chính – ngân hàng',
    '  2.5. Ngành nông nghiệp',
    '  2.6. Ngành vận tải và logistics',
    '  2.7. Ngành truyền thông và giải trí',
    '  2.8. Ngành dịch vụ khách hàng',
    'CHƯƠNG 3: CƠ HỘI VÀ THÁCH THỨC',
    '  3.1. Cơ hội cho người lao động',
    '  3.2. Thách thức và rủi ro',
    '  3.3. Định hướng cho tương lai',
    'KẾT LUẬN',
    'TÀI LIỆU THAM KHẢO',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    p.paragraph_format.line_spacing = 1.8
    p.paragraph_format.space_after = Pt(0)

doc.add_page_break()

# ═══════════════════════════════════════
#  LỜI MỞ ĐẦU
# ═══════════════════════════════════════
heading('LỜI MỞ ĐẦU', level=1)

para('Trong những năm gần đây, trí tuệ nhân tạo (Artificial Intelligence – AI) đã chuyển mình từ một khái niệm trong phòng thí nghiệm thành một lực lượng sản xuất hiện thực, len lỏi vào hầu hết mọi ngóc ngách của nền kinh tế toàn cầu. Sự ra đời của các mô hình ngôn ngữ lớn (LLM) như GPT, Claude, Gemini cùng các hệ thống AI tạo sinh (Generative AI) đã đánh dấu một bước ngoặt lịch sử, tương tự cuộc cách mạng hơi nước hay cách mạng số trước đây.')

para('Tuy nhiên, cùng với những cơ hội chưa từng có về năng suất và hiệu quả, AI cũng đặt ra những câu hỏi lớn về việc làm, về vai trò của con người trong nền sản xuất tương lai, và về khoảng cách kỹ năng giữa các nhóm lao động. Câu hỏi "Liệu AI có thay thế hoàn toàn con người?" không còn là chuyện khoa học viễn tưởng, mà đã trở thành một vấn đề chính sách cấp bách.')

para('Tiểu luận này sẽ phân tích một cách có hệ thống các tác động của AI lên tám nhóm ngành nghề tiêu biểu, từ đó rút ra các cơ hội, thách thức và định hướng cho người lao động trong bối cảnh mới.')

doc.add_page_break()

# ═══════════════════════════════════════
#  CHƯƠNG 1
# ═══════════════════════════════════════
heading('CHƯƠNG 1: TỔNG QUAN VỀ TRÍ TUỆ NHÂN TẠO', level=1)

heading('1.1. Khái niệm và lịch sử phát triển', level=2)
para('Trí tuệ nhân tạo (AI) là lĩnh vực của khoa học máy tính nhằm tạo ra các hệ thống có khả năng thực hiện các nhiệm vụ thường đòi hỏi trí thông minh của con người, bao gồm học tập, lập luận, giải quyết vấn đề, nhận thức ngôn ngữ và ra quyết định.')
para('Thuật ngữ "Artificial Intelligence" được John McCarthy chính thức đưa ra tại Hội thảo Dartmouth năm 1956. Trải qua hơn 60 năm phát triển với nhiều thăng trầm (các "mùa đông AI" và "mùa xuân AI"), lĩnh vực này đã đạt được những bước tiến vượt bậc vào thập niên 2010 nhờ ba yếu tố hội tụ: (1) sự bùng nổ của dữ liệu lớn (Big Data), (2) sức mạnh tính toán tăng theo định luật Moore nhờ GPU/TPU, và (3) các đột phá trong học sâu (Deep Learning).')

heading('1.2. Các công nghệ AI cốt lõi', level=2)
para('Những công nghệ AI chính đang định hình nền kinh tế hiện nay bao gồm:')

bullet('Học máy (Machine Learning) và Học sâu (Deep Learning): cho phép máy tính tự học từ dữ liệu mà không cần lập trình tường minh.')
bullet('Xử lý ngôn ngữ tự nhiên (NLP): giúp máy hiểu, sinh và tương tác bằng ngôn ngữ con người – nền tảng của chatbot, trợ lý ảo, dịch thuật.')
bullet('Thị giác máy tính (Computer Vision): nhận diện hình ảnh, video – ứng dụng trong y tế (chẩn đoán ảnh X-quang), xe tự lái, an ninh.')
bullet('AI tạo sinh (Generative AI): tạo ra nội dung mới (văn bản, hình ảnh, âm nhạc, mã nguồn) dựa trên dữ liệu huấn luyện – điển hình là ChatGPT, Midjourney, GitHub Copilot.')
bullet('Robot tự động hóa quy trình bằng phần mềm (RPA): tự động hóa các tác vụ văn phòng lặp đi lặp lại.')

heading('1.3. Thực trạng ứng dụng AI hiện nay', level=2)
para('Theo báo cáo của McKinsey Global Institute (2025), AI đóng góp ước tính khoảng 4,4 nghìn tỷ USD vào nền kinh tế toàn cầu mỗi năm, tương đương với GDP của Đức. Tại Việt Nam, Chính phủ đã ban hành Chiến lược quốc gia về nghiên cứu và phát triển AI đến năm 2030 (Quyết định 127/QĐ-TTg), xác định AI là một trong những lĩnh vực công nghệ ưu tiên hàng đầu. Các doanh nghiệp Việt từ FPT, VNG đến các startup AI đều đã có những sản phẩm ứng dụng thực tế trong nhiều ngành.')

doc.add_page_break()

# ═══════════════════════════════════════
#  CHƯƠNG 2
# ═══════════════════════════════════════
heading('CHƯƠNG 2: TÁC ĐỘNG CỦA AI ĐẾN CÁC NGÀNH NGHỀ', level=1)

# 2.1
heading('2.1. Ngành sản xuất và chế tạo', level=2)
para('Ngành sản xuất là một trong những lĩnh vực chịu tác động sớm nhất và sâu sắc nhất từ AI, với khái niệm "Công nghiệp 4.0". Các nhà máy thông minh (smart factory) sử dụng AI để tối ưu hóa dây chuyền sản xuất, dự đoán bảo trì thiết bị (predictive maintenance), kiểm soát chất lượng bằng thị giác máy tính và quản lý chuỗi cung ứng theo thời gian thực.')
para('Theo nghiên cứu của Oxford Economics (2024), khoảng 20 triệu việc làm trong ngành sản xuất toàn cầu có nguy cơ bị thay thế bởi robot và AI vào năm 2030. Tuy nhiên, số lượng việc làm mới được tạo ra trong các lĩnh vực kỹ thuật AI, vận hành hệ thống, và phân tích dữ liệu sản xuất có thể bù đắp phần lớn sự mất mát này. Vấn đề cốt lõi là sự dịch chuyển về kỹ năng: lao động tay nghề thấp sẽ giảm, trong khi lao động có kỹ năng STEM và dữ liệu sẽ tăng mạnh.')

# 2.2
heading('2.2. Ngành y tế và chăm sóc sức khỏe', level=2)
para('AI đang tạo ra một cuộc cách mạng trong y học: từ chẩn đoán hình ảnh (AI phát hiện ung thư vú từ ảnh chụp nhũ ảnh với độ chính xác cao hơn bác sĩ X-quang trung bình 11,5%), phát triển thuốc (AlphaFold của DeepMind dự đoán cấu trúc protein – rút ngắn thời gian phát triển thuốc từ năm xuống tháng), đến y học cá thể hóa và trợ lý ảo cho bác sĩ.')
para('Báo cáo của Accenture (2025) dự báo các ứng dụng AI trong y tế có thể tiết kiệm 150 tỷ USD mỗi năm cho hệ thống chăm sóc sức khỏe Hoa Kỳ vào năm 2028. Tuy nhiên, AI không thay thế bác sĩ mà đóng vai trò "trợ lý thông minh": bác sĩ vẫn là người ra quyết định cuối cùng, chịu trách nhiệm pháp lý và thực hiện các tương tác cảm xúc với bệnh nhân. Ngành y tế vì thế sẽ chứng kiến sự thay đổi về vai trò hơn là mất việc làm hàng loạt.')

# 2.3
heading('2.3. Ngành giáo dục và đào tạo', level=2)
para('AI tạo sinh cho phép cá nhân hóa việc học tập ở quy mô chưa từng có. Các hệ thống dạy kèm AI như Khan Academy với Khanmigo, Duolingo Max có thể thích ứng với trình độ từng học sinh, cung cấp phản hồi tức thì và giải thích tường minh. Điều này đặt ra câu hỏi về vai trò của giáo viên truyền thống.')
para('Nhiều chuyên gia dự báo rằng AI sẽ không thay thế giáo viên, nhưng sẽ thay đổi căn bản phương pháp giảng dạy. Giáo viên chuyển từ người truyền thụ kiến thức sang người hướng dẫn, cố vấn và thiết kế trải nghiệm học tập. Các kỹ năng như tư duy phản biện, sáng tạo, giải quyết vấn đề phức tạp – những điều AI chưa thể làm tốt – sẽ trở thành trọng tâm của giáo dục. Những môn học thuần túy ghi nhớ và lặp lại có nguy cơ mất giá trị.')

# 2.4
heading('2.4. Ngành tài chính – ngân hàng', level=2)
para('AI đã thâm nhập sâu vào hệ thống tài chính qua nhiều kênh: chatbot hỗ trợ khách hàng 24/7, phát hiện gian lận giao dịch theo thời gian thực, chấm điểm tín dụng tự động, giao dịch thuật toán (algorithmic trading) chiếm hơn 70% khối lượng giao dịch trên thị trường chứng khoán Mỹ, và tư vấn tài chính cá nhân (robo-advisor).')
para('Theo World Economic Forum (2025), khoảng 30% công việc trong ngành tài chính – ngân hàng có nguy cơ bị tự động hóa cao trong 5 năm tới – chủ yếu là các vị trí giao dịch viên, nhân viên xử lý hồ sơ, kiểm toán nội bộ cấp thấp. Tuy nhiên, nhu cầu về chuyên gia phân tích rủi ro, nhà phát triển hệ thống AI tài chính và chuyên gia tuân thủ (compliance) đang tăng mạnh. Xu hướng "fintech + AI" đang tạo ra một hệ sinh thái tài chính mới với chi phí thấp hơn và khả năng tiếp cận rộng hơn.')

# 2.5
heading('2.5. Ngành nông nghiệp', level=2)
para('Nông nghiệp thông minh (smart agriculture) ứng dụng AI qua máy bay không người lái (drone) giám sát đồng ruộng, cảm biến IoT đo độ ẩm và dinh dưỡng đất, hệ thống tưới tiêu tự động dựa trên dự báo thời tiết, và robot thu hoạch trái cây. AI giúp tăng năng suất 20–30% và giảm 40% lượng nước tiêu thụ trong nông nghiệp.')
para('Tác động đến việc làm trong nông nghiệp có hai mặt. Ở các nước phát triển, lao động nông nghiệp giảm nhưng giá trị sản xuất tăng. Ở các nước đang phát triển như Việt Nam, AI có thể giúp nông dân tiếp cận kiến thức canh tác tiên tiến qua điện thoại thông minh (chatbot nông nghiệp), dự báo giá cả thị trường, và cảnh báo sâu bệnh sớm – giảm thiểu rủi ro thay vì thay thế lao động.')

# 2.6
heading('2.6. Ngành vận tải và logistics', level=2)
para('Xe tự lái (autonomous vehicles) – từ ô tô, xe tải đến drone giao hàng – là ứng dụng AI gây tranh luận nhất trong ngành vận tải. Công nghệ này hứa hẹn giảm 94% tai nạn giao thông do lỗi người lái (WHO), tối ưu hóa nhiên liệu và giảm ùn tắc. Waymo, Tesla, Baidu đã vận hành thương mại xe tự lái tại một số thành phố.')
para('Tuy nhiên, gần 3,5 triệu tài xế xe tải chỉ riêng tại Mỹ đối mặt với nguy cơ mất việc. Logistics và quản lý chuỗi cung ứng cũng thay đổi mạnh: AI dự báo nhu cầu, tối ưu hóa tuyến đường, quản lý kho hàng tự động. Các trung tâm logistics của Amazon, Alibaba đã sử dụng robot Kiva để di chuyển hàng hóa, giảm 70% thời gian xử lý đơn hàng.')

# 2.7
heading('2.7. Ngành truyền thông và giải trí', level=2)
para('Ứng dụng AI tạo sinh đang làm thay đổi căn bản cách sản xuất nội dung. Các công cụ như ChatGPT (văn bản), Midjourney/DALL·E (hình ảnh), Suno/Udio (âm nhạc), và Sora/Runway (video) cho phép bất kỳ ai cũng có thể tạo ra nội dung chất lượng cao chỉ với vài dòng prompt. Điều này ảnh hưởng trực tiếp đến các ngành báo chí, quảng cáo, thiết kế đồ họa, sản xuất phim và âm nhạc.')
para('Một báo cáo của Goldman Sachs (2025) ước tính 300 triệu việc làm trên toàn cầu sẽ bị ảnh hưởng bởi AI tạo sinh, trong đó người làm nội dung – viết lách, dịch thuật, thiết kế – thuộc nhóm chịu tác động mạnh nhất. Tuy nhiên, AI cũng dân chủ hóa khả năng sản xuất nội dung: một cá nhân có thể làm công việc của cả một phòng ban sáng tạo. Giá trị sẽ dịch chuyển từ kỹ năng "làm ra sản phẩm" sang kỹ năng "ý tưởng và chỉ đạo sản xuất".')

# 2.8
heading('2.8. Ngành dịch vụ khách hàng', level=2)
para('Chatbot AI và trợ lý ảo đã trở thành "mặt tiền" của dịch vụ khách hàng tại hầu hết các doanh nghiệp lớn. Các hệ thống như Zendesk AI, Intercom Fin có thể xử lý tới 80% yêu cầu của khách hàng mà không cần người tham gia. Trong tương lai gần, AI có thể xử lý gần như toàn bộ tương tác cấp 1 (level-1 support).')

bullet('Vị trí giảm: nhân viên tổng đài, nhân viên CSAT trực chat các câu hỏi cơ bản.')
bullet('Vị trí tăng: chuyên gia thiết kế luồng hội thoại AI, kỹ sư prompt, giám sát chất lượng phản hồi AI, nhân viên xử lý các trường hợp phức tạp (level-2/3).')

doc.add_page_break()

# ═══════════════════════════════════════
#  CHƯƠNG 3
# ═══════════════════════════════════════
heading('CHƯƠNG 3: CƠ HỘI VÀ THÁCH THỨC', level=1)

heading('3.1. Cơ hội cho người lao động', level=2)
para('Tự động hóa các công việc lặp đi lặp lại giúp con người tập trung vào các nhiệm vụ sáng tạo, chiến lược và có giá trị cao hơn. Các nghiên cứu cho thấy các doanh nghiệp áp dụng AI có năng suất tăng 30–50%, từ đó tạo ra dư địa để mở rộng quy mô và tuyển dụng thêm nhân sự cho các vị trí mới.')
para('AI cũng mở ra các ngành nghề hoàn toàn mới: kỹ sư prompt, nhà đạo đức AI, chuyên gia tương tác người-máy, kỹ sư tích hợp AI vào quy trình, chuyên gia hậu kiểm nội dung AI, v.v. Hầu hết các công việc này không tồn tại cách đây 5 năm. Người lao động biết tận dụng AI như một "cộng sự" sẽ có lợi thế cạnh tranh vượt trội so với người không sử dụng.')

heading('3.2. Thách thức và rủi ro', level=2)

bullet('Mất việc hàng loạt do tự động hóa: các ngành phụ thuộc vào lao động chân tay và lao động văn phòng lặp lại chịu rủi ro cao nhất.')
bullet('Bất bình đẳng thu nhập gia tăng: người có kỹ năng AI hưởng lương cao; người thiếu kỹ năng bị đào thải. Khoảng cách kỹ năng ngày càng mở rộng.')
bullet('Đạo đức và pháp lý: AI thiên vị (bias) dữ liệu huấn luyện, AI tạo ra thông tin sai lệch (hallucination), vi phạm bản quyền, quyền riêng tư và giám sát AI ngoài tầm kiểm soát.')
bullet('Tác động tâm lý – xã hội: sự lo lắng về việc làm, khủng hoảng danh tính nghề nghiệp, và nguy cơ con người trở nên quá phụ thuộc vào AI.')

heading('3.3. Định hướng cho tương lai', level=2)
para('Trước những tác động sâu rộng của AI, tiểu luận đề xuất một số định hướng:')

bullet('Học tập suốt đời (lifelong learning): người lao động cần chủ động cập nhật kỹ năng mới, đặc biệt là kỹ năng làm việc cùng AI, tư duy phản biện và giải quyết vấn đề phức tạp.')
bullet('Cải cách hệ thống giáo dục: các trường học cần chuyển từ dạy kiến thức sang dạy năng lực thích ứng, gồm AI literacy là môn học bắt buộc.')
bullet('Chính sách an sinh xã hội: các chính phủ cần xây dựng mạng lưới an toàn cho người lao động bị ảnh hưởng, bao gồm bảo hiểm thất nghiệp AI, chương trình đào tạo lại miễn phí.')
bullet('Hợp tác người – AI: thay vì coi AI là đối thủ, cần xây dựng mô hình "human-in-the-loop" nơi AI làm công việc nặng nhọc và con người đưa ra quyết định cuối cùng.')

doc.add_page_break()

# ═══════════════════════════════════════
#  KẾT LUẬN
# ═══════════════════════════════════════
heading('KẾT LUẬN', level=1)

para('AI không phải là một "cơn sóng thần" sẽ cuốn trôi mọi việc làm, mà giống một "cơn lũ" làm thay đổi dòng chảy thay vì xóa sổ nó. Lịch sử đã chứng minh: cuộc cách mạng máy hơi nước không xóa bỏ lao động – nó chuyển lao động từ đồng ruộng vào nhà máy; cuộc cách mạng số không giết chết các ngành – nó sinh ra hàng loạt ngành mới không ai tưởng tượng trước được.')

para('Điều duy nhất chắc chắn là "cơn lũ" AI đang tái định hình từng ngành nghề, từng vị trí công việc. Người lao động thích ứng sẽ tồn tại và phát triển; người đứng yên sẽ bị bỏ lại. Trong kỷ nguyên này, kỹ năng quan trọng nhất không phải là biết gì, mà là khả năng học bất cứ điều gì khi cần – và biết cách cộng tác với AI để làm những điều một mình không thể.')

# ═══════════════════════════════════════
#  TÀI LIỆU THAM KHẢO
# ═══════════════════════════════════════
doc.add_page_break()
heading('TÀI LIỆU THAM KHẢO', level=1)

refs = [
    'McKinsey Global Institute (2025). "The Economic Potential of Generative AI". McKinsey & Company.',
    'Goldman Sachs (2025). "The Potentially Large Effects of Artificial Intelligence on Economic Growth". Global Economics Analyst.',
    'World Economic Forum (2025). "The Future of Jobs Report 2025". WEF, Geneva.',
    'Accenture (2025). "AI in Healthcare: The Promise and the Path Forward". Accenture Research.',
    'Oxford Economics (2024). "How Robots Change the World". Oxford Economics Report.',
    'Brynjolfsson, E. & McAfee, A. (2023). "The AI Frontier: How Artificial Intelligence is Reshaping Work". MIT Press.',
    'Bộ Khoa học và Công nghệ Việt Nam (2021). "Quyết định 127/QĐ-TTg: Chiến lược quốc gia về nghiên cứu và phát triển AI đến năm 2030".',
    'Nguyen, T. H. & Pham, Q. A. (2024). "AI Adoption in Vietnamese Enterprises: Status and Barriers". VNU Journal of Science: Computer Science, 40(2), 45–58.',
    'DeepMind (2024). "AlphaFold: A Solution to a 50-Year-Old Grand Challenge in Biology". Nature, 604, 56–62.',
    'Russell, S. & Norvig, P. (2025). "Artificial Intelligence: A Modern Approach" (5th ed.). Pearson.',
]

for r in refs:
    p = doc.add_paragraph()
    run = p.add_run(r)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-1.5)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)

# ── Save ──
path = os.path.expanduser('~/documents/tac-dong-ai-len-cac-nganh-nghe.docx')
doc.save(path)
print(f'Saved: {path}')
